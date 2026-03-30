"""
File loader: reads CSV, Excel, PDF, DOCX and returns a pandas DataFrame.
"""
import io
import pandas as pd
import pdfplumber
from docx import Document


def load_file(file_path: str, file_type: str) -> pd.DataFrame:
    """
    Load a file from disk into a DataFrame.
    file_type: 'csv' | 'xlsx' | 'xls' | 'pdf' | 'docx'
    """
    ft = file_type.lower().lstrip('.')

    if ft == 'csv':
        return _load_csv(file_path)
    elif ft in ('xlsx', 'xls'):
        return _load_excel(file_path)
    elif ft == 'pdf':
        return _load_pdf(file_path)
    elif ft == 'docx':
        return _load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _load_csv(file_path: str) -> pd.DataFrame:
    encodings = ['utf-8', 'latin-1', 'cp1252']
    for enc in encodings:
        try:
            df = pd.read_csv(file_path, encoding=enc)
            if not df.empty:
                return df
        except (UnicodeDecodeError, pd.errors.ParserError):
            continue
    raise ValueError("Could not decode CSV file with any common encoding.")


def _load_excel(file_path: str) -> pd.DataFrame:
    xl = pd.ExcelFile(file_path)
    frames = []
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        if not df.empty:
            df['_sheet'] = sheet
            frames.append(df)
    if not frames:
        raise ValueError("No data found in Excel file.")
    combined = pd.concat(frames, ignore_index=True)
    return combined


def _load_pdf(file_path: str) -> pd.DataFrame:
    frames = []
    text_rows = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            # Try to extract tables first
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    if table and len(table) > 1:
                        header = [str(c).strip() if c else f'col_{i}' for i, c in enumerate(table[0])]
                        for row in table[1:]:
                            row_dict = {header[i]: (str(cell).strip() if cell else '') for i, cell in enumerate(row)}
                            row_dict['_page'] = page_num
                            frames.append(row_dict)
            else:
                # Fall back to text extraction per paragraph
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        line = line.strip()
                        if line:
                            text_rows.append({'text_content': line, '_page': page_num})

    if frames:
        return pd.DataFrame(frames)
    elif text_rows:
        return pd.DataFrame(text_rows)
    else:
        raise ValueError("No extractable data found in PDF.")


def _load_docx(file_path: str) -> pd.DataFrame:
    doc = Document(file_path)
    frames = []
    text_rows = []

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        # Deduplicate header columns
        seen = {}
        clean_header = []
        for col in header:
            if col in seen:
                seen[col] += 1
                clean_header.append(f"{col}_{seen[col]}")
            else:
                seen[col] = 0
                clean_header.append(col)

        for row in table.rows[1:]:
            row_dict = {clean_header[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
            row_dict['_table'] = table_idx
            frames.append(row_dict)

    # Extract paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_rows.append({'text_content': text})

    if frames:
        return pd.DataFrame(frames)
    elif text_rows:
        return pd.DataFrame(text_rows)
    else:
        raise ValueError("No extractable data found in DOCX.")
